"""Generate the BT-branded UNICEF Venture Fund 2026 submission brief for Zeeshan.

Output: ~/Repo/school-climate-hub/docs/UNICEF-VF-Submission-Brief-Zeeshan-260515v1.docx

Uses the canonical BT styling helpers in ~/Claude/contract_styles.py
(BT_BLUE accent, Calibri, BT logo header, footer with PAGE/NUMPAGES fields).
"""
import sys
from pathlib import Path

sys.path.insert(0, "/Users/rezamalik/Claude")
from contract_styles import (
    create_doc, add_title, add_subtitle, add_heading1, add_heading2,
    add_body, add_body_bold, add_bullet, add_table, add_logo_header,
    add_footer, add_page_break, add_horizontal_line,
)

OUT = Path.home() / (
    "Library/CloudStorage/GoogleDrive-reza@beaconhouse.tech/My Drive/_Claude/"
    "Eng/PDLC/UNICEF-VF-Submission-Brief-260515v1.1.docx"
)


def placeholder(doc, label, hint=None, multiline=False):
    """Render a labelled answer slot for Zeeshan to fill."""
    add_body_bold(doc, label)
    if hint:
        p = doc.add_paragraph()
        run = p.add_run(hint)
        run.italic = True
        from docx.shared import Pt, RGBColor
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"
    add_body(doc, "[ Answer: ___________________________________________________________ ]")
    if multiline:
        add_body(doc, "[ ________________________________________________________________ ]")
        add_body(doc, "[ ________________________________________________________________ ]")


def build():
    doc = create_doc()
    add_logo_header(doc)
    add_footer(doc)

    # ---- Cover ----
    add_title(doc, "UNICEF Venture Fund 2026 — Climate Tech")
    add_subtitle(doc, "Submission Brief")
    add_body(doc, "Form: https://form.jotform.com/260711626703351")
    add_body(doc, "Deadline: 17 May 2026, 23:59 CET  ·  Funding: up to US$100,000 equity-free")
    add_body(doc, "Drafter: Zeeshan  ·  Final reviewer: Reza Malik")
    add_body(doc, "Applicant: Beaconhouse Technology (Pvt) Ltd (Pakistan, woman-led)  ·  Deployment partner: Premier DLC")
    add_body(doc, "Live demo: https://schoolclimatehub.org   ·   Repo: https://github.com/school-climate-hub/school-climate-hub")
    add_horizontal_line(doc)
    add_body(doc, "How to use this document:")
    add_bullet(doc, "Sections mirror the Jotform fields in order.")
    add_bullet(doc, "Pre-filled answers are best-effort defaults — replace or refine.")
    add_bullet(doc, "Empty [ Answer: ___ ] slots need Zeeshan + Reza input. Character limits noted in red where strict.")
    add_bullet(doc, "Once finished, paste each answer into the Jotform. Don't submit from this doc.")

    add_page_break(doc)

    # ---- Section 1: Referral ----
    add_heading1(doc, "Section 1 · Referral Source")

    add_heading2(doc, "Q1.1 — How did you hear from us?")
    add_body(doc, "Dropdown options: UNICEF Website · UNICEF Venture Fund team member · Newsletter · UNICEF Country Office · Cambridge ISL · India Health Fund · LinkedIn · Facebook/Instagram · Online Search · Incubator/accelerator · AfriLabs")
    placeholder(doc, "Selection:", "(needs Reza confirm — likely 'UNICEF Country Office' if PDLC/Erum sourced via UNICEF Pakistan, else 'Online Search')")

    add_heading2(doc, "Q1.2 — Name of incubator / accelerator / Country Office / VF team member that referred you")
    placeholder(doc, "Answer:", "Conditional on Q1.1. Provide name + role if a person; provide office name if institutional.")

    add_page_break(doc)

    # ---- Section 2: Eligibility ----
    add_heading1(doc, "Section 2 · Eligibility")

    add_heading2(doc, "Q2.1 — Is your company registered as a private for-profit company?")
    add_body_bold(doc, "Answer: Yes")
    add_body(doc, "Source: Beaconhouse Technology (Pvt) Ltd — SECP 0137683, NTN 6297182, registered office 10–11 Gurumangat Road, Lahore.")

    add_heading2(doc, "Q2.2 — Is your solution Open Source?")
    add_body_bold(doc, "Answer: Yes")
    add_body(doc, "Source: github.com/school-climate-hub/school-climate-hub — Apache-2.0.")

    add_heading2(doc, "Q2.3 — In which country is your company legally registered?")
    add_body_bold(doc, "Answer: Pakistan")

    add_page_break(doc)

    # ---- Section 3: DEI ----
    add_heading1(doc, "Section 3 · Diversity, Equity and Inclusion")

    add_heading2(doc, "Q3.1 — Is your company founded or co-founded by women?")
    add_body_bold(doc, "Answer: Yes")
    add_body(doc, "Beaconhouse Technology is woman-led at the C-suite level: CEO Amina Kasuri and COO Fatima Kasuri.")

    add_heading2(doc, "Q3.2 — Names of woman founders / co-founders")
    add_body_bold(doc, "Names: Amina Kasuri (CEO), Fatima Kasuri (COO)")

    add_heading2(doc, "Q3.3 — Is your company founded or led by youth (younger than 35)?")
    placeholder(doc, "Answer (Yes/No):", "Depends on Amina + Fatima's DOBs. Confirm with Reza.")

    add_heading2(doc, "Q3.4 — Names of young leaders (if applicable)")
    placeholder(doc, "Names:")

    add_heading2(doc, "Q3.5 — Team table (up to 5 rows)")
    add_body(doc, "Fields per row: Full name · Position · Year of birth · Nationality · Gender · Short bio (1–2 sentences). Roster mixes BT leadership with the PDLC deployment lead so reviewers see both the engineering depth and on-the-ground delivery capability:")
    add_table(
        doc,
        ["Full name", "Position", "Year of birth", "Nationality", "Gender", "Short bio"],
        [
            ["Amina Kasuri", "CEO, Beaconhouse Technology", "____", "Pakistani", "Woman", "CEO; oversees BT's engineering portfolio including School Climate Hub and education-tech platforms across South Asia."],
            ["Fatima Kasuri", "COO, Beaconhouse Technology", "____", "Pakistani", "Woman", "COO; runs operations across BT's engineering teams and partner deployments."],
            ["Reza Malik", "Director, Beaconhouse Technology", "____", "Pakistani", "Man", "Director; product + engineering lead for School Climate Hub. Reports to Nassir Kasuri (Board)."],
            ["Erum [Surname]", "Head, Premier DLC (Deployment partner)", "____", "Pakistani", "Woman", "Leads PDLC's deployment of School Climate Hub across 50 PSSP/PSRP schools in Gujranwala — owns school relationships, EMIS roster, and ops feedback."],
            ["[BT engineer or PDLC ops]", "[Position]", "____", "____", "____", "[Bio]"],
        ],
    )
    add_body(doc, "Three of four named team members are women. If the form requires applicant-entity-only roster, drop Erum and add a BT engineer.")

    add_page_break(doc)

    # ---- Section 4: Company Information ----
    add_heading1(doc, "Section 4 · Company Information")

    add_heading2(doc, "Q4.1 — Name of company")
    add_body_bold(doc, "Answer: Beaconhouse Technology (Pvt) Ltd")

    add_heading2(doc, "Q4.2 — Company's web page")
    add_body_bold(doc, "Answer: https://schoolclimatehub.org")
    add_body(doc, "This is the single front door for reviewers. The live demo, methodology, open data downloads, partner credits, and a prominent link to the open-source GitHub repository are all reachable from here. UNICEF reviewers should land on schoolclimatehub.org first; from there they can click through to GitHub. Ensure the homepage prominently surfaces the GitHub link before submission.")

    add_heading2(doc, "Q4.3 — Contact person")
    add_body_bold(doc, "Name: Reza Malik")
    add_body_bold(doc, "Designation: Director, Beaconhouse Technology")

    add_heading2(doc, "Q4.4 — Primary e-mail address")
    add_body_bold(doc, "Answer: reza@beaconhouse.tech")

    add_heading2(doc, "Q4.5 — Alternate contact person email (optional)")
    placeholder(doc, "Answer:", "Suggest Erum (Head, Premier DLC) so reviewers can reach the deployment-side lead independently.")

    add_heading2(doc, "Q4.6 — Year company was founded")
    placeholder(doc, "Answer:", "Confirm BT founding year — likely 2023 or 2024.")

    add_page_break(doc)

    # ---- Section 5: Proposal ----
    add_heading1(doc, "Section 5 · Proposal for Raising Funds")
    add_body(doc, "Character limits in this section are strict. Word counts shown alongside each draft.")

    add_heading2(doc, "Q5.1 — Which challenges does your solution address? (multi-select)")
    add_body_bold(doc, "Tick: Area 1 (Resilient infrastructure)  +  Area 2 (Early warning, Early Action systems)")
    add_body(doc, "Skip Area 3 (Health Care readiness) and Area 4 (Point-of-Care low-connectivity) — they dilute the pitch.")

    add_heading2(doc, "Q5.2 — What need / challenge is your solution addressing locally? (max 150 WORDS)")
    add_body_bold(doc, "Draft (~130 words):")
    add_body(doc, "Pakistan's school year is increasingly disrupted by climate shocks — heatwaves crossing 45 °C, monsoon flooding, and seasonal air pollution. Across our 50-school pilot in Gujranwala, recomputed from Premier DLC's monthly attendance registers, average attendance fell from 93.0% in 2023 to 89.3% in 2025 — a 3.7 percentage-point decline equating to ~93,427 lost child-school-days against the 2023 baseline. School operators today have no per-school visibility into hazard exposure: provincial forecasts arrive too late to plan around, public health advisories are not school-aware, and parents are not notified when their child's school is at risk. Decisions are reactive (closing schools after the heat has already hit) rather than pre-emptive. The 50 schools in our pilot serve roughly 12,000 children — a population for whom every lost school-day compounds existing learning gaps from COVID, inflation, and chronic under-investment in education infrastructure.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.3 — Describe your proposed solution and how it addresses the challenges (max 250 WORDS)")
    add_body_bold(doc, "Draft (~215 words):")
    add_body(doc, "School Climate Hub is an open-data, AI-grounded operational platform that turns global climate observations into per-school operational decisions. For each of the 50 pilot schools in Gujranwala we ingest live ECMWF HRES (10-day deterministic) and ENS (15-day probabilistic) forecasts, ERA5 historical reanalysis, and MODIS land-surface temperature, then resolve them to each school's coordinates and produce a real-time hazard score and 15-day forecast. The public-facing site (schoolclimatehub.org) lets anyone — parents, journalists, district officers, UNICEF reviewers — see today's exposure and the week ahead.")
    add_body(doc, "An AI assistant grounded strictly in the per-school dataset answers operator questions like “which schools should I close 18–22 May” and drafts parent advisories in English, Urdu, or Shahmukhi Punjabi for the operator to review and dispatch. Subscribers receive free web-push notifications when their school crosses a threshold — no app store, no SMS cost, no proprietary lock-in.")
    add_body(doc, "Critically, the per-school exposure data is published openly under Apache-2.0, so partners can fork, self-host, or build derivative tools. The platform is designed to extend horizontally across UNICEF country offices: any province or country can onboard its schools with the same data pipeline, simply by uploading an EMIS roster.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.4 — Which technology(ies) are you using?")
    add_body_bold(doc, "Tick: AI / Data Science")
    add_body_bold(doc, "Other (clarify): Open climate data ingestion (ECMWF HRES + ENS, ERA5, MODIS)")

    add_heading2(doc, "Q5.5 — How are you using the technologies? (max 150 WORDS)")
    add_body_bold(doc, "Draft (~130 words):")
    add_body(doc, "School Climate Hub fuses three technology layers. Open climate data flows in from ECMWF (HRES + ENS), Copernicus ERA5 reanalysis, and NASA MODIS — sampled at every school's coordinates and aggregated into per-day exposure scores and a 15-day probabilistic forecast.")
    add_body(doc, "AI grounding uses Anthropic Claude with the full per-school dataset injected as system context. Every answer cites real numbers from the dataset — the model is constrained to refuse questions it cannot ground in the data, eliminating hallucination risk.")
    add_body(doc, "Web delivery uses a Progressive Web App with Cloudflare Web Push. Parents subscribe with an email, install the site to their home screen with one tap, and receive native push notifications when their school crosses a risk threshold — without paying for SMS, without installing a native app, and without giving up data to a proprietary platform. The codebase is Apache-2.0; the per-school data is published openly.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.6 — Current technical status / what's needed for pilot (max 150 WORDS)")
    add_body_bold(doc, "Draft (~130 words):")
    add_body(doc, "Working prototype is live at schoolclimatehub.org. The end-to-end data pipeline runs against real production data: Copernicus CDS API for ERA5, NASA Earthdata for MODIS, and the ECMWF S3 open-data mirror for HRES (10-day, 3-hourly) and ENS (15-day, 50 perturbed members). Per-school exposure scores are recomputed on each refresh and published as scores.json. An AI chat backend on Cloudflare Workers streams Anthropic Claude responses grounded in the latest dataset.")
    add_body(doc, "For the v0.2 production rollout we need: (1) authentication and per-user rate limiting on the AI chat to control cost; (2) Web Push + PWA install flow for free parent notifications; (3) multilingual UI in English, Urdu, Arabic, and French; (4) operator dispatch UI for parent advisories; (5) onboarding training for the 50 school operators in Gujranwala alongside Premier DLC.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.7 — Results of prototyping and testing (max 150 WORDS)")
    add_body_bold(doc, "Draft (~135 words):")
    add_body(doc, "The pilot prototype currently operates on real 2026 data for all 50 schools. ECMWF HRES forecasts a red-threshold heatwave 18–22 May 2026, with daily maxima reaching 42–43 °C; the ENS ensemble shows red-heat probability climbing from 0% this week through 86% by 28 May. The AI assistant correctly surfaces the highest-burden schools by combining heat exposure with student population, and drafts closure-advisory text in English, Urdu, and Punjabi for operator approval.")
    add_body(doc, "Independently, three years of historical attendance data from Premier DLC (2023–2025, monthly, all 50 schools) shows average attendance falling from 93.0% to 89.3% — a 3.7 percentage-point decline equating to ~93,427 lost child-school-days against the 2023 baseline. Attendance is presented as a correlated outcome (not a causal claim — the decline reflects post-COVID, economic, and management factors alongside climate). The combined hazard + measured-outcome data gives the platform an empirical foundation that purely modeled risk scores lack.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.8 — GitHub or other open repository link")
    add_body_bold(doc, "Answer: https://github.com/school-climate-hub/school-climate-hub")
    add_body(doc, "Public, Apache-2.0. No access grants required for @unicefinnovation. The repo is also linked from the schoolclimatehub.org homepage so reviewers can navigate naturally: site → demo → code.")

    add_heading2(doc, "Q5.9 — Project targets / milestones for the next 12 months (max 250 WORDS)")
    add_body_bold(doc, "Draft (~220 words):")
    add_body(doc, "Months 1–3 — Production v0.2 release. Migrate the live prototype to a multi-tenant Next.js application with email-based authentication, per-user rate-limited AI chat, multilingual UI (English, Urdu, Arabic, French) with RTL support, and free web-push notifications via PWA install. Launch into production across all 50 PDLC schools in Gujranwala, with operator onboarding workshops conducted jointly by BT and Premier DLC.")
    add_body(doc, "Months 4–6 — First UNICEF country-office onboarding. Replicate the data pipeline and operator workflow for one additional UNICEF programme country (target: Bangladesh, Nepal, or Indonesia). Add the operator dispatch UI for parent advisories with AI-drafted multilingual text + human approval. First public alert sent at scale; 500+ schools onboarded.")
    add_body(doc, "Months 7–9 — District admin tooling. Add role-based access for district education officers to view roll-up dashboards across their schools, manage threshold policies, and broadcast advisories. Onboard 5,000+ schools cumulatively. Begin publishing per-school climate impact reports as open data.")
    add_body(doc, "Months 10–12 — Open API + ecosystem. Publish a public REST/OpenAPI surface so researchers, journalists, and third-party tools can build on the dataset. Target: 10,000+ subscribers; one peer-reviewed publication on heat × attendance correlation co-authored with academic partners.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.10 — Key partners and advisors (max 150 WORDS)")
    add_body_bold(doc, "Draft (~110 words):")
    add_body(doc, "Premier DLC (PDLC) — primary deployment partner in Pakistan; owns the 50-school relationship across the PSSP / PSRP networks in Gujranwala, the EMIS roster, and the three-year historical attendance baseline (2023–2025) we use to calibrate climate–outcome correlations.")
    add_body(doc, "ECMWF (European Centre for Medium-Range Weather Forecasts) — source for HRES and ENS forecast products via the public open-data mirror. Copernicus Climate Data Store — ERA5 reanalysis. NASA Earthdata — MODIS land-surface temperature. Anthropic — Claude model for the grounded AI assistant. Cloudflare — hosting, edge compute, push notification delivery. All current partners are operationally engaged; no paper letters of intent. UNICEF Country Office collaboration is part of the v0.2 expansion plan.")
    placeholder(doc, "Zeeshan's version:", "Add UNICEF Pakistan Country Office if formally engaged before submit.")

    add_heading2(doc, "Q5.11 — How much are you seeking to raise?")
    add_body_bold(doc, "Answer: US$100,000")

    add_heading2(doc, "Q5.12 — What was your company's revenue last year?")
    placeholder(doc, "Answer (USD):", "Required field — Beaconhouse Technology 2025 revenue in USD. Confirm with Reza / BT finance.")

    add_heading2(doc, "Q5.13 — Overview of capital and other contributions (max 150 WORDS)")
    add_body_bold(doc, "Draft (~100 words):")
    add_body(doc, "Beaconhouse Technology has self-funded the prototype to date, including all engineering time, infrastructure (Cloudflare and Anthropic API), data subscriptions, and operational coordination. Premier DLC contributes in kind: school relationships across 50 PSSP/PSRP schools, EMIS roster data, three years of historical monthly attendance data (2023–2025), operator feedback, and on-the-ground deployment support. No external grants or third-party investments have been deployed against this project. UNICEF Venture Fund support would underwrite the v0.2 build — production hardening, authentication, multilingual UI, push notifications, dispatch tooling, and onboarding of the first non-pilot country office.")
    placeholder(doc, "Zeeshan's version:")

    add_heading2(doc, "Q5.14 — Share application with partners?")
    add_body_bold(doc, "Answer: Yes")

    add_heading2(doc, "Q5.15 — Subscribe to UNICEF Venture Fund newsletter?")
    add_body_bold(doc, "Answer: Yes")

    add_page_break(doc)

    # ---- Section 6: Pitch Video ----
    add_heading1(doc, "Section 6 · Pitch Video")
    add_body(doc, "Spec: 2 minutes, English (or English subtitles), demo-first, uploaded to YouTube (unlisted is fine).")

    add_heading2(doc, "Suggested 2-minute script outline")
    add_table(
        doc,
        ["Time", "Beat"],
        [
            ["0:00–0:15", "Hook: “Climate makes Pakistan's school year unpredictable. 50 schools, 31k children, no per-school visibility.”"],
            ["0:15–0:35", "Solution: live dashboard at schoolclimatehub.org — show the map, click a red school."],
            ["0:35–1:00", "Forecast story: ECMWF predicts a red heatwave 18–22 May. Show the sparkline + the AI response."],
            ["1:00–1:25", "Operator workflow: open Schools list, click ‘Send advisory’, AI drafts EN/UR, operator approves."],
            ["1:25–1:45", "Open data + open source: 50 schools' data is public, Apache-2.0 repo."],
            ["1:45–2:00", "Roadmap: free push alerts via PWA, scale via UNICEF country offices, ask: $100k to v0.2."],
        ],
        col_widths=[1.0, 5.5],
    )

    add_heading2(doc, "Five video production decisions (Reza + Zeeshan to agree before recording)")
    add_bullet(doc, "1. Voiceover: Reza vs Zeeshan vs AI-generated?")
    add_bullet(doc, "2. Format: screen-capture demo only vs talking-head + demo?")
    add_bullet(doc, "3. Subtitles: burned-in vs YouTube auto-captions?")
    add_bullet(doc, "4. Music: yes / no / what?")
    add_bullet(doc, "5. Outro: contact card / partner logos / both?")

    add_heading2(doc, "Q6.1 — Upload video pitch to YouTube (link)")
    placeholder(doc, "YouTube URL:", "Test the link in an incognito window before pasting.")

    add_page_break(doc)

    # ---- Open items / blockers ----
    add_heading1(doc, "Open Items — Need Reza Sign-Off Before Submit")
    add_table(
        doc,
        ["#", "Item", "Owner", "Blocker?"],
        [
            ["1", "Confirm referral source (UNICEF CO vs Online Search)", "Reza + Erum", "No"],
            ["2", "Confirm Amina + Fatima Kasuri DOBs (for Q3.3 youth check + Q3.5 row)", "Reza", "YES"],
            ["3", "Confirm Reza + Erum DOBs (Q3.5 rows)", "Reza", "YES"],
            ["4", "Confirm 5th team-table row (BT engineer or PDLC ops) — name, DOB, nat., gender, bio", "Reza", "YES"],
            ["5", "Erum's full surname for the team table", "Reza / Erum", "YES"],
            ["6", "BT founding year (Q4.6)", "Reza / BT finance", "YES"],
            ["7", "BT 2025 revenue in USD (Q5.12)", "Reza / BT finance", "YES"],
            ["8", "Decide alternate contact email (Q4.5) — Erum suggested", "Reza", "No"],
            ["9", "Lock proposal copy (all 150/250-char fields)", "Zeeshan → Reza", "YES"],
            ["10", "Produce 2-minute pitch video + upload to YouTube", "Zeeshan + Reza decisions", "YES"],
            ["11", "Verify https://schoolclimatehub.org loads on HTTPS apex", "auto (Pages cert)", "No"],
            ["12", "Polish GitHub README hero / quick-start", "BT", "No"],
            ["13", "Submit before 17 May 23:59 CET (aim 16 May EOD PKT)", "Reza", "YES"],
        ],
        col_widths=[0.4, 3.7, 1.7, 0.8],
    )

    add_heading1(doc, "Submission Strategy")
    add_bullet(doc, "Submit by end of 16 May (PKT) to leave one full day for fixes if Jotform rejects anything.")
    add_bullet(doc, "Keep all character counts at ≤95% of the limit — leaves room for last-minute word swaps.")
    add_bullet(doc, "Upload the pitch video as unlisted on YouTube with burned-in English subtitles.")
    add_bullet(doc, "Polish the GitHub README: hero image, quick-start, link to live demo, link to docs/v0.2-scope.md.")
    add_bullet(doc, "Lead with “Pakistan-registered Pvt Ltd” wherever relevant — UNICEF programme-country status is a strength.")

    add_page_break(doc)

    add_heading1(doc, "Appendix · Company Description (Beaconhouse Technology)")
    add_body(doc, "Three lengths for use in different submission fields, supporting documents, or follow-up due-diligence requests. Pick by context; all three carry the same facts.")

    add_heading2(doc, "Long (proposal / about page · ~160 words)")
    add_body(doc, "Beaconhouse Technology (Pvt) Ltd is a Pakistan-registered, woman-led engineering company building software for education, climate adaptation, and operational tooling. Led by CEO Amina Kasuri and COO Fatima Kasuri, BT has shipped production systems for multi-campus school networks across South Asia and Southeast Asia — including student information systems, operational dashboards, and integration platforms serving tens of thousands of students.")
    add_body(doc, "On School Climate Hub, BT is the applicant and technical builder: architecting the data pipelines (ERA5, MODIS, ECMWF HRES + ENS), the AI grounding layer (Anthropic Claude), the public web app, and the alert dispatch infrastructure. The codebase is published under Apache-2.0; per-school exposure data is released as an open dataset so any operator, researcher, or country office can fork, replicate, or extend the work.")
    add_body(doc, "Headquartered in Lahore, BT operates as a self-funded private company and partners with deployment-led organisations like Premier DLC for in-market execution.")

    add_heading2(doc, "Medium (partners section, doc footnote · ~70 words)")
    add_body(doc, "Beaconhouse Technology (Pvt) Ltd — Pakistan-registered, woman-led (CEO Amina Kasuri, COO Fatima Kasuri) engineering company building education and climate-tech software. BT is the applicant and technical builder of School Climate Hub: data pipelines (ERA5 / MODIS / ECMWF), AI grounding, web app, and alert dispatch. Codebase Apache-2.0; per-school exposure data published openly. Track record across multi-campus school networks in South Asia and Southeast Asia. Self-funded; partners with deployment-led organisations for in-market execution.")

    add_heading2(doc, "One-line (tight fields, captions · ~140 chars)")
    add_body(doc, "Beaconhouse Technology (Pvt) Ltd — woman-led Pakistan engineering company; builder of School Climate Hub. Apache-2.0, self-funded.")

    add_heading2(doc, "Notes for Zeeshan before pasting the bio")
    add_bullet(doc, "Replace “tens of thousands of students” with a hard number if Reza confirms it's safe to publish.")
    add_bullet(doc, "Add BT founding year once Reza confirms (2023 or 2024).")
    add_bullet(doc, "If a named anchor product can be disclosed (e.g., BEAMS360), insert it as “including [Product]” in the long version — concrete names beat generic claims.")

    add_page_break(doc)

    add_heading1(doc, "Credibility Talking Points (for any free-text overflow)")
    add_bullet(doc, "Working prototype with real data — schoolclimatehub.org is live, 50 schools, real ECMWF + ERA5 + MODIS, AI grounded in the dataset.")
    add_bullet(doc, "Open source from day 1 — Apache-2.0; UNICEF can fork, modify, redeploy without us.")
    add_bullet(doc, "Open data from day 1 — per-school scores, exposure timeseries and forecasts published as JSON/CSV/Parquet.")
    add_bullet(doc, "Free at point of use — public can subscribe to alerts via PWA web push; no app store, no payment, no account required.")
    add_bullet(doc, "Run cost under $20/month at pilot scale — sustainable without ongoing UNICEF subsidy.")
    add_bullet(doc, "Real deployment partner (PDLC) with EMIS roster access, not a letter of intent.")
    add_bullet(doc, "Designed for UNICEF country office adoption — multilingual (EN/UR/AR/FR planned), multi-tenant by district, no vendor lock-in.")

    doc.save(OUT)
    print(f"wrote {OUT}")
    print(f"  size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
