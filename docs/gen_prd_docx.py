"""Generate the BT-branded PRD.docx from docs/PRD.md.

Uses the canonical ~/Claude/contract_styles.py helpers (BT_BLUE, Calibri,
logo header, footer with PAGE/NUMPAGES fields). Markdown-driven so future
PRD updates regenerate by re-running the script.

Outputs:
  - docs/PRD.docx  (in-repo copy)
  - _Claude/Eng/PDLC/BT-Eng-School-Climate-Hub-PRD-260516v2.0.docx  (Drive)

Supported markdown subset:
  - # / ## / ### headings
  - Paragraphs (with **bold** and *italic* inline)
  - Bullet lists (- or *)
  - Numbered lists (1. 2. ...)
  - Pipe tables (| col | col |)
  - Fenced code blocks (``` ... ```)
  - Inline `code`
  - Links [text](url) — rendered as styled text
  - Horizontal rule (---)
  - Checkboxes ([ ] / [x])
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, "/Users/rezamalik/Claude")
from contract_styles import (
    create_doc, add_title, add_heading1, add_heading2,
    add_table, add_logo_header, add_footer, add_horizontal_line,
    BT_BLUE, BT_DARK, BT_GRAY,
)
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/rezamalik/Repo/school-climate-hub")
PRD_MD = ROOT / "docs" / "PRD.md"
PRD_DOCX = ROOT / "docs" / "PRD.docx"
DRIVE_COPY = Path(
    "/Users/rezamalik/Library/CloudStorage/GoogleDrive-reza@beaconhouse.tech/"
    "My Drive/_Claude/Eng/PDLC/BT-Eng-School-Climate-Hub-PRD-260516v2.0.docx"
)


def add_heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BT_BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = BT_DARK
    # Light grey background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)


def emit_inline(paragraph, text):
    """Emit inline text into an existing paragraph with **bold**, *italic*, `code`, [text](url)."""
    # Tokenise into segments
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**") and chunk.endswith("**"):
            r = paragraph.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith("*") and chunk.endswith("*"):
            r = paragraph.add_run(chunk[1:-1]); r.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = paragraph.add_run(chunk[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif chunk.startswith("[") and "](" in chunk:
            label_end = chunk.index("](")
            label = chunk[1:label_end]
            r = paragraph.add_run(label)
            r.font.color.rgb = BT_BLUE
            r.font.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, i):
    """Parse a markdown pipe table starting at line i. Returns (headers, rows, next_i)."""
    def split_row(line):
        parts = [p.strip() for p in line.strip().strip("|").split("|")]
        return parts

    headers = split_row(lines[i])
    # Skip separator row (|---|---|)
    if i + 1 < len(lines) and re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]):
        i += 2
    else:
        i += 1
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1
    return headers, rows, i


def render_markdown(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", stripped):
            add_horizontal_line(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading1(doc, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading2(doc, stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading3(doc, stripped[5:].strip())
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|?[\s\-:|]+\|?\s*$", lines[i + 1]
        ):
            headers, rows, i = parse_table(lines, i)
            add_table(doc, headers, rows)
            continue

        # Bulleted lists
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line).rstrip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            emit_inline(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line).rstrip()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            emit_inline(p, text)
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Italic-only paragraph (e.g., the closing footer-style line)
        if stripped.startswith("*") and stripped.endswith("*") and "**" not in stripped:
            p = doc.add_paragraph()
            r = p.add_run(stripped[1:-1])
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = BT_GRAY
            p.paragraph_format.space_before = Pt(8)
            i += 1
            continue

        # Default: paragraph with inline formatting
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        emit_inline(p, stripped)
        i += 1


def build():
    doc = create_doc()
    add_logo_header(doc)
    add_footer(doc)

    md_text = PRD_MD.read_text(encoding="utf-8")
    render_markdown(doc, md_text)

    PRD_DOCX.parent.mkdir(parents=True, exist_ok=True)
    DRIVE_COPY.parent.mkdir(parents=True, exist_ok=True)
    doc.save(PRD_DOCX)
    doc.save(DRIVE_COPY)
    print(f"wrote {PRD_DOCX} ({PRD_DOCX.stat().st_size / 1024:.1f} KB)")
    print(f"wrote {DRIVE_COPY} ({DRIVE_COPY.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    build()
